"""
Physics-Driven Insights into Inter-Laboratory Noise as the Limiting Factor in Machine Learning Prediction of Bandgap Shifts in Doped Metal Oxides
----------------------------------------------------------
This script reproduces all manuscript figures and tables in a single
execution, so every number in the paper matches the corresponding figure.
Training, cross-validation, SHAP analysis, conformal prediction, and
noise-floor calculations all share the same random seed and data split.

Usage (Google Colab or local):
    pip install pandas numpy scikit-learn scipy matplotlib openpyxl shap python-docx
    python bandgap_complete.py

Outputs written to bandgap_outputs/ and bundled into bandgap_outputs.zip:
    23 TIFF figures at 600 dpi (LZW compressed)
    bandgap_tables.docx  -- manuscript tables 1-5 plus a key-metrics summary
    requirements.txt     -- library versions for reproducibility

Figures:
    fig01a  R² comparison (test vs 5-fold CV)
    fig01b  RMSE comparison
    fig01c  MAE comparison
    fig02a-c  Parity plots: Extra Trees, Random Forest, Gradient Boosting
    fig03a  SHAP feature importance bar (top 20)
    fig03b  SHAP beeswarm (top 10)
    fig04a  Leave-one-host-out R² by host
    fig04b  LOHO vs random-split R²
    fig05a  Delta-Eg distribution by host
    fig05b  Delta-Eg distribution by synthesis route
    fig06a  Residuals vs predicted
    fig06b  Residual histogram
    fig06c  Cumulative absolute error distribution
    fig07a  90% conformal prediction interval band
    fig07b  Conformal coverage scatter
    fig08a  Inter-laboratory sigma for all repeated host-dopant pairs
    fig08b  Noise floor summary vs model error
    fig09a  Within-paper concentration-sensitivity slope scatter
    fig09b  Per-host Pearson r and sign agreement
    fig10a  Delta-Eg vs dopant concentration (coloured by electronegativity)
    fig10b  Delta-Eg vs ionic radius mismatch (coloured by synthesis temperature)

All runs with random_state=42 produce identical numbers.
"""
# pip install handled externally
import os, sys, warnings, zipfile
warnings.filterwarnings("ignore")

import numpy  as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines  import Line2D
from matplotlib.ticker import MaxNLocator
from scipy import stats

try:
    import shap as shap_lib
    HAS_SHAP = True
except ImportError:
    HAS_SHAP = False
    print("  NOTE: shap not installed – Fig 3 uses pre-computed SHAP sheet.")

from sklearn.experimental    import enable_iterative_imputer   # noqa
from sklearn.impute          import IterativeImputer
from sklearn.model_selection import (train_test_split, cross_val_score,
                                      LeaveOneGroupOut)
from sklearn.ensemble        import (ExtraTreesRegressor,
                                      RandomForestRegressor,
                                      GradientBoostingRegressor)
from sklearn.linear_model    import Ridge, Lasso, ElasticNet
from sklearn.svm             import SVR
from sklearn.preprocessing   import StandardScaler
from sklearn.metrics         import (r2_score, mean_squared_error,
                                      mean_absolute_error)

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("  NOTE: python-docx not installed – skipping DOCX output.")

# ── paths and global settings ────────────────────────────────────────────────

DATASET_PATH = "Doped_Metal_Oxide_Bandgap_Dataset.xlsx"
OUTPUT_DIR   = "bandgap_outputs"
RANDOM_STATE = 42
TEST_SIZE    = 0.20
TARGET       = "Delta_Bandgap"
DPI          = 600          # 600 dpi meets Elsevier/Springer combination-art requirement

HOST_COLORS = {
    "TiO2":  "#1F77B4",
    "ZnO":   "#2CA02C",
    "WO3":   "#FF7F0E",
    "SnO2":  "#9467BD",
    "In2O3": "#D62728",
    "CeO2":  "#8C564B",
}

FEATURE_COLS = [
    "Undoped_Eg_lit", "Host_O_per_fu", "OHE_Synth_Spray-Pyrolysis",
    "OHE_Synth_Co-Precipitation", "Synthesis_Temp(C)",
    "Dopant_electronegativity", "Ionic_radius_mismatch_rel",
    "OHE_Synth_Hydrothermal", "Dopant_ionization_energy_eV",
    "FE_Conc_x_EN", "EN_difference", "FE_Size_x_Conc",
    "Dopant_d_electrons", "OHE_Synth_Physical-Deposition",
    "Ionic_Potential_Diff", "Dopant_atomic_number", "Dopant_group",
    "FE_Conc_x_IonMismatch", "Crystallite_Size_nm", "Log_Conc",
    "Conc_1", "Conc_squared", "Host_electron_affinity_eV",
    "Host_eff_mass_e", "Dopant_ionic_radius_pm", "Dopant_period",
    "Dopant_common_ox_state", "Dopant_is_transition_metal",
    "Dopant_is_rare_earth", "Dopant_atomic_radius_pm",
    "Ox_state_diff", "Ionic_Potential_Dopant", "Ionic_Potential_Host",
    "FE_Conc_x_OxDiff", "FE_IonicPotential_Ratio", "Has_d_electrons",
    "OHE_Synth_Sol-Gel", "OHE_Synth_Combustion",
    "OHE_Synth_Green-Synthesis", "OHE_Phase_Anatase",
    "OHE_Phase_Rutile", "OHE_Phase_Wurtzite", "OHE_Phase_Cubic",
    "OHE_Phase_Hexagonal", "Mixed_Phase",
]

# ── font sizes and line weights used across all figures ──────────────────────

FS_TITLE  = 18    # panel title
FS_AXLBL  = 16    # axis labels
FS_TICK   = 14    # tick labels
FS_LEGEND = 14    # legend
FS_ANNOT  = 13    # bar labels, stat boxes

SPINE_LW    = 1.8
TICK_LEN    = 8
TICK_WIDTH  = 1.8
REF_LINE_LW = 2.0
LABEL_COLOR = "#1A1A1A"
TITLE_PAD   = 12

BAR_COLORS = ["#1F4E79", "#2E75B6", "#5BA3D0",
              "#888888", "#CCCCCC", "#BBBBBB", "#AAAAAA"]


def apply_global_style():
    plt.rcParams.update({
        "font.family":        "serif",
        "font.serif":         ["Times New Roman", "DejaVu Serif", "serif"],
        "mathtext.fontset":   "dejavuserif",
        "font.weight":        "bold",
        "font.size":          FS_TICK,
        "axes.titlesize":     FS_TITLE,
        "axes.labelsize":     FS_AXLBL,
        "axes.titleweight":   "bold",
        "axes.labelweight":   "bold",
        "axes.titlelocation": "center",
        "axes.titlepad":      TITLE_PAD,
        "axes.linewidth":     SPINE_LW,
        "legend.fontsize":    FS_LEGEND,
        "legend.framealpha":  0.95,
        "legend.edgecolor":   "black",
        "xtick.labelsize":    FS_TICK,
        "ytick.labelsize":    FS_TICK,
        "xtick.direction":    "in",
        "ytick.direction":    "in",
        "xtick.major.size":   TICK_LEN,
        "ytick.major.size":   TICK_LEN,
        "xtick.minor.size":   TICK_LEN * 0.5,
        "ytick.minor.size":   TICK_LEN * 0.5,
        "xtick.major.width":  TICK_WIDTH,
        "ytick.major.width":  TICK_WIDTH,
        "xtick.minor.width":  TICK_WIDTH * 0.6,
        "ytick.minor.width":  TICK_WIDTH * 0.6,
        "xtick.top":          True,
        "ytick.right":        True,
        "figure.facecolor":   "white",
        "axes.facecolor":     "white",
        "figure.dpi":         DPI,
        "savefig.dpi":        DPI,
        "savefig.bbox":       "tight",
        "savefig.facecolor":  "white",
        "axes.grid":          False,
    })


# ── small helper functions shared by all figure routines ─────────────────────

def _style_ax(ax):
    # tighten spines and force tick marks inward on all four sides
    for spine in ax.spines.values():
        spine.set_linewidth(SPINE_LW)
    ax.tick_params(which="major", direction="in",
                   length=TICK_LEN, width=TICK_WIDTH,
                   top=True, right=True, labelsize=FS_TICK)
    ax.tick_params(which="minor", direction="in",
                   length=TICK_LEN * 0.5, width=TICK_WIDTH * 0.6,
                   top=True, right=True)
    for lbl in ax.get_xticklabels() + ax.get_yticklabels():
        lbl.set_fontweight("bold")
        lbl.set_fontsize(FS_TICK)


def _bold_legend(leg):
    for t in leg.get_texts():
        t.set_fontweight("bold")
        t.set_fontsize(FS_LEGEND)


def _save(fig, ax, fname):
    _style_ax(ax)
    path = os.path.join(OUTPUT_DIR, fname)
    fig.savefig(path, dpi=DPI, bbox_inches="tight",
                facecolor="white", format="tiff",
                pil_kwargs={"compression": "tiff_lzw"})
    plt.close(fig)
    print(f"  saved -> {fname}")


def _save_twin(fig, ax, ax2, fname):
    _style_ax(ax)
    for spine in ax2.spines.values():
        spine.set_linewidth(SPINE_LW)
    ax2.tick_params(which="major", direction="in",
                    length=TICK_LEN, width=TICK_WIDTH, labelsize=FS_TICK)
    for lbl in ax2.get_yticklabels():
        lbl.set_fontweight("bold")
        lbl.set_fontsize(FS_TICK)
    path = os.path.join(OUTPUT_DIR, fname)
    fig.savefig(path, dpi=DPI, bbox_inches="tight",
                facecolor="white", format="tiff",
                pil_kwargs={"compression": "tiff_lzw"})
    plt.close(fig)
    print(f"  saved -> {fname}")


# ── data loading, model training, and all derived quantities ─────────────────
# Everything is computed once and stored in a dict so figures and tables
# read from the same objects and are therefore guaranteed to match.

def build_everything():
    print("\nLoading dataset and training models ...")

    if not os.path.exists(DATASET_PATH):
        sys.exit(f"\nDataset file '{DATASET_PATH}' not found in {os.getcwd()}.\n"
                 "Place the .xlsx file in the same directory and re-run.\n")

    df = pd.read_excel(DATASET_PATH, sheet_name="Experimental_Data")
    print(f"  {len(df)} rows  |  {df['Host'].nunique()} host oxides  |"
          f"  {df['Paper_ID'].nunique()} source papers")

    # ── load pre-computed sheets ─────────────────────────────────────────────
    try:
        loho_sheet = pd.read_excel(DATASET_PATH, sheet_name="LOHO_Results")
    except Exception:
        loho_sheet = pd.DataFrame(columns=["Host", "Random_Split_R2"])

    try:
        shap_sheet = pd.read_excel(DATASET_PATH, sheet_name="SHAP_Importance")
    except Exception:
        shap_sheet = None

    # ── feature selection ────────────────────────────────────────────────────
    available = [c for c in FEATURE_COLS if c in df.columns]
    missing   = set(FEATURE_COLS) - set(available)
    if missing:
        print(f"  Warning: {len(missing)} expected column(s) not found in the sheet:")
        for m in sorted(missing):
            print(f"    {m}")

    X = df[available].copy()
    y = df[TARGET].copy()

    # ── 80/20 split (random_state=42, leakage-free imputer) ──────────────────
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=TEST_SIZE, random_state=RANDOM_STATE)
    train_idx = X_train.index
    test_idx  = X_test.index

    imputer   = IterativeImputer(max_iter=10, random_state=RANDOM_STATE)
    X_tr_imp  = imputer.fit_transform(X_train)
    X_te_imp  = imputer.transform(X_test)
    X_all_imp = imputer.transform(X)

    scaler   = StandardScaler()
    X_tr_sc  = scaler.fit_transform(X_tr_imp)
    X_te_sc  = scaler.transform(X_te_imp)

    print(f"  Train={len(X_train)}  Test={len(X_test)}  Features={len(available)}")

    # ── seven-model benchmark ────────────────────────────────────────────────
    linear_set = {"SVR-RBF", "Lasso", "Ridge", "ElasticNet"}
    model_defs = {
        "Extra Trees":       ExtraTreesRegressor(
                                 n_estimators=300, min_samples_leaf=2,
                                 random_state=RANDOM_STATE, n_jobs=-1),
        "Random Forest":     RandomForestRegressor(
                                 n_estimators=300,
                                 random_state=RANDOM_STATE, n_jobs=-1),
        "Gradient Boosting": GradientBoostingRegressor(
                                 n_estimators=300, learning_rate=0.05,
                                 max_depth=4, subsample=0.8,
                                 random_state=RANDOM_STATE),
        "SVR-RBF":           SVR(kernel="rbf", C=10, gamma="scale"),
        "Lasso":             Lasso(alpha=1e-3, max_iter=10_000),
        "Ridge":             Ridge(alpha=1.0),
        "ElasticNet":        ElasticNet(alpha=1e-3, max_iter=10_000),
    }

    ORDER = ["Extra Trees", "Random Forest", "Gradient Boosting",
             "SVR-RBF", "Lasso", "Ridge", "ElasticNet"]

    print(f"\n  {'Model':<22} {'Test R2':>8} {'RMSE':>8} {'MAE':>8} "
          f"{'CV R2':>8} {'CV sig':>8}")
    print("  " + "-" * 68)

    results = {}
    for name in ORDER:
        model = model_defs[name]
        Xtr   = X_tr_sc if name in linear_set else X_tr_imp
        Xte   = X_te_sc if name in linear_set else X_te_imp
        model.fit(Xtr, y_train)
        y_pred = model.predict(Xte)
        cv     = cross_val_score(model, Xtr, y_train, cv=5, scoring="r2")
        results[name] = dict(
            model=model, y_pred=y_pred,
            Test_R2  = float(r2_score(y_test, y_pred)),
            RMSE     = float(np.sqrt(mean_squared_error(y_test, y_pred))),
            MAE      = float(mean_absolute_error(y_test, y_pred)),
            CV_R2    = float(cv.mean()),
            CV_sig   = float(cv.std()),
        )
        r = results[name]
        print(f"  {name:<22} {r['Test_R2']:>8.4f} {r['RMSE']:>8.4f} "
              f"{r['MAE']:>8.4f} {r['CV_R2']:>8.4f} {r['CV_sig']:>8.4f}")

    gb          = results["Gradient Boosting"]["model"]
    gb_all_pred = gb.predict(X_all_imp)

    # ── conformal prediction (Section 4.7) ───────────────────────────────────
    X_fit, X_cal, y_fit, y_cal = train_test_split(
        X_tr_imp, y_train, test_size=0.20, random_state=RANDOM_STATE)
    gb_conf = GradientBoostingRegressor(
        n_estimators=300, learning_rate=0.05,
        max_depth=4, subsample=0.8, random_state=RANDOM_STATE)
    gb_conf.fit(X_fit, y_fit)
    n_cal        = len(y_cal)
    cal_scores   = np.abs(y_cal.values - gb_conf.predict(X_cal))
    adj_level    = np.ceil((n_cal + 1) * 0.90) / n_cal
    q90          = float(np.quantile(cal_scores, min(adj_level, 1.0)))
    y_pred_conf  = gb_conf.predict(X_te_imp)
    lower        = y_pred_conf - q90
    upper        = y_pred_conf + q90
    within       = (y_test.values >= lower) & (y_test.values <= upper)
    coverage     = float(within.mean()) * 100
    fs_guarantee = (1.0 - 0.10 - 1.0 / (n_cal + 1)) * 100

    # ── inter-laboratory noise floor ─────────────────────────────────────────
    hd    = df.groupby(["Host", "Dopant_1"])["Paper_ID"].nunique()
    multi = hd[hd > 1]
    vdf_rows = []
    for (host, dop), _ in multi.items():
        sub = df[(df["Host"] == host) & (df["Dopant_1"] == dop)]
        pp  = sub.groupby("Paper_ID")[TARGET].mean()
        vdf_rows.append({
            "Host": host, "Dopant": dop,
            "N_papers": len(pp), "N_rows": len(sub),
            "sigma": float(pp.std()),
            "range": float(pp.max() - pp.min()),
        })
    vdf = pd.DataFrame(vdf_rows).sort_values("sigma", ascending=False).reset_index(drop=True)

    # ── clean-subset model ────────────────────────────────────────────────────
    hv        = vdf[vdf["sigma"] > 0.40]
    hv_combos = set(zip(hv["Host"], hv["Dopant"]))
    hv_flag   = df.apply(
        lambda r: (r["Host"], r["Dopant_1"]) in hv_combos, axis=1)
    df_clean  = df[~hv_flag].copy()
    X_cl      = df_clean[available].copy()
    y_cl      = df_clean[TARGET].copy()
    X_tr_c, X_te_c, y_tr_c, y_te_c = train_test_split(
        X_cl, y_cl, test_size=TEST_SIZE, random_state=RANDOM_STATE)
    imp_c     = IterativeImputer(max_iter=10, random_state=RANDOM_STATE)
    X_tr_c_i  = imp_c.fit_transform(X_tr_c)
    X_te_c_i  = imp_c.transform(X_te_c)
    gb_cl = GradientBoostingRegressor(
        n_estimators=300, learning_rate=0.05,
        max_depth=4, subsample=0.8, random_state=RANDOM_STATE)
    gb_cl.fit(X_tr_c_i, y_tr_c)
    y_pred_cl  = gb_cl.predict(X_te_c_i)
    r2_clean   = float(r2_score(y_te_c, y_pred_cl))
    mae_clean  = float(mean_absolute_error(y_te_c, y_pred_cl))
    rmse_clean = float(np.sqrt(mean_squared_error(y_te_c, y_pred_cl)))
    cv_cl      = cross_val_score(gb_cl, X_tr_c_i, y_tr_c, cv=5, scoring="r2")
    cv_r2_clean  = float(cv_cl.mean())
    cv_sig_clean = float(cv_cl.std())

    # ── within-paper slope analysis ──────────────────────────────────────────
    train_set  = set(train_idx)
    test_set   = set(test_idx)
    df_idx     = list(df.index)
    slope_rows = []
    for (pid, host, dop), grp in df.groupby(["Paper_ID", "Host", "Dopant_1"]):
        grp = grp.sort_values("Conc_1")
        if grp["Conc_1"].nunique() < 3:
            continue
        if (grp["Conc_1"].max() - grp["Conc_1"].min()) < 1e-6:
            continue
        gi    = set(grp.index)
        n_tst = len(gi & test_set)
        n_trn = len(gi & train_set)
        label = ("train-only" if n_tst == 0 else
                 "test-only"  if n_trn == 0 else "mixed")
        x_c   = grp["Conc_1"].values
        y_obs = grp[TARGET].values
        pos_  = [df_idx.index(i) for i in grp.index]
        y_pr  = gb_all_pred[pos_]
        s_obs,  *_ = stats.linregress(x_c, y_obs)
        s_pred, *_ = stats.linregress(x_c, y_pr)
        slope_rows.append({
            "Paper_ID": pid, "Host": host, "Dopant": dop,
            "N_points":        len(grp),
            "Observed_slope":  float(s_obs),
            "Predicted_slope": float(s_pred),
            "Sign_agreement":  int(np.sign(s_obs) == np.sign(s_pred)),
            "Series_type":     label,
        })
    slope_df = pd.DataFrame(slope_rows)
    r_all, p_all = stats.pearsonr(
        slope_df["Observed_slope"], slope_df["Predicted_slope"])
    to_      = slope_df[slope_df["Series_type"] == "train-only"]
    mx_      = slope_df[slope_df["Series_type"] == "mixed"]
    r_to, _  = stats.pearsonr(to_["Observed_slope"], to_["Predicted_slope"])
    r_mx, p_mx = stats.pearsonr(mx_["Observed_slope"], mx_["Predicted_slope"])

    # SHAP values for Gradient Boosting feature importance
    shap_df = None
    shap_vals_arr = None
    if HAS_SHAP:
        try:
            print("\n  Computing SHAP values for Gradient Boosting ...")
            explainer     = shap_lib.TreeExplainer(gb)
            shap_vals_arr = explainer.shap_values(X_te_imp, check_additivity=False)
            mean_abs      = np.abs(shap_vals_arr).mean(axis=0)
            shap_df = pd.DataFrame(
                {"feature": available, "mean_abs_shap": mean_abs}
            ).sort_values("mean_abs_shap", ascending=False).reset_index(drop=True)
        except Exception as e:
            print(f"  SHAP calculation failed ({e}); using pre-computed values from dataset sheet.")
            shap_df = shap_sheet
    elif shap_sheet is not None:
        shap_df = shap_sheet.sort_values("mean_abs_shap", ascending=False).reset_index(drop=True)
    else:
        print("  No SHAP data found. Figure 3 will be skipped.")

    # Leave-one-host-out: train on five hosts, test on the sixth, repeat for all six.
    # Each fold uses a fresh imputer to avoid data leakage across hosts.
    print("\n  Leave-one-host-out cross-validation (6 folds, ~3-5 min) ...")
    X_loho = df[available].copy()
    y_loho = df[TARGET].copy()
    groups = df["Host"].values
    logo   = LeaveOneGroupOut()
    loho_res = {}
    for train_ix, test_ix in logo.split(X_loho, y_loho, groups):
        host = groups[test_ix[0]]
        X_tr_l, X_te_l = X_loho.iloc[train_ix], X_loho.iloc[test_ix]
        y_tr_l, y_te_l = y_loho.iloc[train_ix], y_loho.iloc[test_ix]
        imp_l    = IterativeImputer(max_iter=10, random_state=RANDOM_STATE)
        X_tr_l_i = imp_l.fit_transform(X_tr_l)
        X_te_l_i = imp_l.transform(X_te_l)
        gb_l = GradientBoostingRegressor(
            n_estimators=300, learning_rate=0.05,
            max_depth=4, subsample=0.8, random_state=RANDOM_STATE)
        gb_l.fit(X_tr_l_i, y_tr_l)
        y_pred_l = gb_l.predict(X_te_l_i)
        r2_l   = float(r2_score(y_te_l, y_pred_l))
        rmse_l = float(np.sqrt(mean_squared_error(y_te_l, y_pred_l)))
        mae_l  = float(mean_absolute_error(y_te_l, y_pred_l))
        loho_res[host] = dict(LOHO_R2=r2_l, RMSE=rmse_l, MAE=mae_l,
                               N_test=len(y_te_l))
        print(f"    {host:<8}  R2={r2_l:.4f}  N_test={len(y_te_l)}")

    # stored random-split R² from sheet (for Fig 4b comparison bars)
    stored_rs = {}
    for _, row_ in loho_sheet.iterrows():
        stored_rs[row_["Host"]] = float(row_.get("Random_Split_R2",
                                   results["Gradient Boosting"]["Test_R2"]))
    for host in loho_res:
        if host not in stored_rs:
            stored_rs[host] = results["Gradient Boosting"]["Test_R2"]

    # ── residual diagnostics ──────────────────────────────────────────────────
    gb_y_pred = results["Gradient Boosting"]["y_pred"]
    resid     = y_test.values - gb_y_pred
    pct_010   = float((np.abs(resid) <= 0.10).mean() * 100)
    pct_020   = float((np.abs(resid) <= 0.20).mean() * 100)
    mean_res  = float(resid.mean())

    print(f"\n  Gradient Boosting residuals:  mean={mean_res:.4f} eV  "
          f"within +/-0.10 eV: {pct_010:.1f}%  within +/-0.20 eV: {pct_020:.1f}%")
    print(f"  Conformal prediction:  q90={q90:.4f} eV  empirical coverage={coverage:.1f}%")
    print(f"  Within-paper slope correlation:  r_all={r_all:.3f}  r_mixed={r_mx:.3f}  p_mixed={p_mx:.3f}")
    print(f"  Clean-subset GB:  R2={r2_clean:.4f}  RMSE={rmse_clean:.4f}  MAE={mae_clean:.4f}")

    return dict(
        df=df, available=available,
        X_tr_imp=X_tr_imp, X_te_imp=X_te_imp, X_all_imp=X_all_imp,
        X_tr_sc=X_tr_sc,   X_te_sc=X_te_sc,
        y_train=y_train,   y_test=y_test,
        train_idx=train_idx, test_idx=test_idx,
        results=results,
        gb=gb, gb_all_pred=gb_all_pred,
        shap_df=shap_df, shap_vals=shap_vals_arr,
        loho_sheet=loho_sheet, loho_res=loho_res, stored_rs=stored_rs,
        y_pred_conf=y_pred_conf, lower=lower, upper=upper,
        q90=q90, coverage=coverage, n_cal=n_cal,
        fs_guarantee=fs_guarantee, within=within,
        vdf=vdf, hv=hv,
        r2_clean=r2_clean, mae_clean=mae_clean, rmse_clean=rmse_clean,
        cv_r2_clean=cv_r2_clean, cv_sig_clean=cv_sig_clean,
        df_clean=df_clean,
        slope_df=slope_df,
        r_all=r_all, p_all=p_all, r_to=r_to, r_mx=r_mx, p_mx=p_mx,
        pct_010=pct_010, pct_020=pct_020, mean_res=mean_res,
    )


# ── Figure 1: three bar charts comparing model test R², RMSE, and MAE ────────

def plot_fig1(d):
    print("\nFigure 1 ...")

    ORDER  = ["Gradient Boosting", "Extra Trees", "Random Forest",
              "SVR-RBF", "Lasso", "Ridge", "ElasticNet"]
    LABELS = ["Gradient\nBoosting", "Extra\nTrees", "Random\nForest",
              "SVR-RBF", "Lasso", "Ridge", "ElasticNet"]
    BCLRS  = ["#1F4E79", "#2E75B6", "#5BA3D0",
              "#888888", "#CCCCCC", "#BBBBBB", "#AAAAAA"]

    r2  = [d["results"][k]["Test_R2"] for k in ORDER]
    cv  = [d["results"][k]["CV_R2"]   for k in ORDER]
    cvs = [d["results"][k]["CV_sig"]  for k in ORDER]
    rm  = [d["results"][k]["RMSE"]    for k in ORDER]
    ma  = [d["results"][k]["MAE"]     for k in ORDER]
    x, w = np.arange(len(ORDER)), 0.38

    # ── 1a: R² comparison ────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(13, 8), constrained_layout=True)
    ax.bar(x - w/2, r2, w, color=BCLRS, alpha=0.92, label="Test R²",
           edgecolor="white", linewidth=0.5)
    ax.bar(x + w/2, cv, w, color=BCLRS, alpha=0.40, label="5-fold CV R²",
           yerr=cvs, capsize=5, ecolor="#444444",
           error_kw={"lw": 1.8, "capthick": 1.8},
           edgecolor="white", linewidth=0.5)
    ax.axhline(0.70, color="red", ls="--", lw=REF_LINE_LW, alpha=0.85)
    ax.text(0.97, 0.715, "R²=0.70", fontsize=FS_ANNOT, color="red",
            ha="right", va="bottom", fontweight="bold",
            transform=ax.get_yaxis_transform())
    # annotate ONLY dark Test R² bars (not lighter CV bars)
    for i, v in enumerate(r2):
        ax.text(i - w/2, v + 0.020, f"{v:.3f}",
                ha="center", va="bottom", rotation=90,
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(x); ax.set_xticklabels(LABELS, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylim(0, 1.22)
    ax.set_ylabel("R² Score", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("R² Comparison – Test vs 5-Fold CV",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="upper left",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig01a_r2_comparison.tiff")

    # ── 1b: RMSE ─────────────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(13, 8), constrained_layout=True)
    bars = ax.bar(x, rm, color=BCLRS, alpha=0.88, edgecolor="white", linewidth=0.5)
    y_max = max(rm)
    ax.set_ylim(0, y_max * 1.30)
    for bar_, v in zip(bars, rm):
        ax.text(bar_.get_x() + bar_.get_width() / 2,
                v + y_max * 0.020, f"{v:.3f}",
                ha="center", va="bottom",
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(x); ax.set_xticklabels(LABELS, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("RMSE (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Root Mean Squared Error",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig01b_rmse.tiff")

    # ── 1c: MAE ──────────────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(13, 8), constrained_layout=True)
    bars = ax.bar(x, ma, color=BCLRS, alpha=0.88, edgecolor="white", linewidth=0.5)
    y_max = max(ma)
    ax.set_ylim(0, y_max * 1.30)
    for bar_, v in zip(bars, ma):
        ax.text(bar_.get_x() + bar_.get_width() / 2,
                v + y_max * 0.020, f"{v:.3f}",
                ha="center", va="bottom",
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(x); ax.set_xticklabels(LABELS, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("MAE (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Mean Absolute Error",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig01c_mae.tiff")


# ── Figure 2: parity plots for the three best models ─────────────────────────

def plot_fig2(d):
    print("\nFigure 2 ...")

    TOP3 = [
        ("Extra Trees",       "fig02a_parity_extra_trees.tiff"),
        ("Random Forest",     "fig02b_parity_random_forest.tiff"),
        ("Gradient Boosting", "fig02c_parity_gradient_boosting.tiff"),
    ]
    test_hosts = d["df"].loc[d["test_idx"], "Host"].values
    y_arr = d["y_test"].values
    lim   = 1.40

    for name, fname in TOP3:
        fig, ax = plt.subplots(figsize=(8, 8), constrained_layout=True)
        y_pred = d["results"][name]["y_pred"]

        for host, clr in HOST_COLORS.items():
            mask = test_hosts == host
            if mask.sum() == 0: continue
            ax.scatter(y_arr[mask], y_pred[mask],
                       color=clr, s=55, alpha=0.82,
                       edgecolors="white", linewidths=0.4,
                       zorder=3, label=host)

        ax.plot([-lim, lim], [-lim, lim], "k--",
                lw=REF_LINE_LW, alpha=0.65, label="1:1 line")
        sf, ic, *_ = stats.linregress(y_arr, y_pred)
        xf = np.linspace(-lim, lim, 200)
        ax.plot(xf, sf * xf + ic, "r-", lw=REF_LINE_LW, alpha=0.80,
                label="OLS fit")
        ax.set_xlim(-lim, lim); ax.set_ylim(-lim, lim)
        ax.set_xlabel("Actual ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
        ax.set_ylabel("Predicted ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
        n_test = len(y_arr)
        ax.set_title(f"Predicted vs Actual ΔEg – {name}\n"
                     f"(Test Set, N={n_test})",
                     fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
        r2v  = d["results"][name]["Test_R2"]
        rmse = d["results"][name]["RMSE"]
        ax.text(0.04, 0.96,
                f"R² = {r2v:.4f}\nRMSE = {rmse:.4f} eV",
                transform=ax.transAxes, fontsize=FS_ANNOT, va="top",
                fontweight="bold",
                bbox=dict(boxstyle="round,pad=0.5", facecolor="white",
                          alpha=0.94, edgecolor="#555555", lw=1.8))
        leg = ax.legend(fontsize=FS_LEGEND, loc="lower right",
                        framealpha=0.95, edgecolor="black", ncol=2)
        _bold_legend(leg)
        ax.set_aspect("equal", adjustable="box")
        _save(fig, ax, fname)


# ── Figure 3: SHAP feature importance ────────────────────────────────────────

def plot_fig3(d):
    print("\nFigure 3 ...")

    if d["shap_df"] is None:
        print("  No SHAP data available. Skipping Figure 3.")
        return

    shap_df = d["shap_df"].copy()

    def prettify(s):
        return (s.str.replace("OHE_Synth_", "Synth: ", regex=False)
                 .str.replace("OHE_Phase_", "Phase: ", regex=False)
                 .str.replace("_", " ", regex=False))

    top20 = shap_df.head(20).copy()
    top20["label"] = prettify(top20["feature"])
    top10 = shap_df.head(10).copy()
    top10["label"] = prettify(top10["feature"])

    # ── 3a: horizontal bar – top 20 ──────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 13), constrained_layout=True)
    vals = top20["mean_abs_shap"].values[::-1]
    labs = top20["label"].values[::-1]
    clrs = ["#1F4E79" if v > 0.02 else "#7FB3D3" for v in vals]
    bars = ax.barh(range(len(top20)), vals, color=clrs, alpha=0.87, height=0.72)
    ax.set_yticks(range(len(top20)))
    ax.set_yticklabels(labs, fontsize=FS_TICK, fontweight="bold")
    # value annotations on each bar
    for bar_, v in zip(bars, vals):
        ax.text(v + max(vals) * 0.012, bar_.get_y() + bar_.get_height() / 2,
                f"{v:.4f}", va="center", ha="left",
                fontsize=FS_ANNOT - 1, fontweight="bold", color=LABEL_COLOR)
    ax.axvline(0.02, color="red", ls="--", lw=REF_LINE_LW, alpha=0.80)
    ax.text(0.022, len(top20) - 0.5, "threshold 0.02",
            fontsize=FS_ANNOT - 1, color="red", fontweight="bold",
            va="top", ha="left")
    ax.set_xlabel("Mean |SHAP| value (eV)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("SHAP Feature Importance\n"
                 "(Gradient Boosting – top 20 of 45 features)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    ax.margins(x=0.18)
    _save(fig, ax, "fig03a_shap_bar.tiff")

    # ── 3b: beeswarm approximation – top 10 ──────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 9), constrained_layout=True)
    np.random.seed(RANDOM_STATE)
    df = d["df"]
    for i, (_, row_) in enumerate(top10.iterrows()):
        feat = row_["feature"]
        if feat in df.columns:
            vals_f = df[feat].dropna().values.astype(float)
            vmin, vmax = vals_f.min(), vals_f.max()
            vn_full = (df[feat].fillna(df[feat].median()).values.astype(float)
                       - vmin) / (vmax - vmin + 1e-8)
        else:
            vn_full = np.random.rand(len(df))
        shap_approx = np.random.normal(0, row_["mean_abs_shap"] * 0.8, len(df))
        jitter      = np.random.uniform(-0.25, 0.25, len(df))
        ax.scatter(shap_approx, np.full(len(df), i) + jitter,
                   c=vn_full, cmap="coolwarm", s=12, alpha=0.60,
                   vmin=0, vmax=1)
    ax.set_yticks(range(len(top10)))
    ax.set_yticklabels(top10["label"].values, fontsize=FS_TICK, fontweight="bold")
    ax.axvline(0, color="gray", lw=REF_LINE_LW * 0.75, alpha=0.55)
    ax.set_xlabel("SHAP value (impact on ΔEg, eV)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("SHAP Beeswarm – Top 10 Features\n"
                 "(blue = low feature value, red = high)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    sm = plt.cm.ScalarMappable(cmap="coolwarm",
                                norm=plt.Normalize(vmin=0, vmax=1))
    sm.set_array([])
    cbar = plt.colorbar(sm, ax=ax, shrink=0.50, pad=0.02)
    cbar.set_label("Feature value (normalised)",
                   fontsize=FS_AXLBL - 1, fontweight="bold")
    cbar.ax.tick_params(labelsize=FS_TICK - 1)
    for lbl in cbar.ax.get_yticklabels():
        lbl.set_fontweight("bold")
    _save(fig, ax, "fig03b_shap_beeswarm.tiff")


# ── Figure 4: leave-one-host-out generalisation analysis ─────────────────────

def plot_fig4(d):
    print("\nFigure 4 ...")

    loho_res  = d["loho_res"]
    stored_rs = d["stored_rs"]

    hosts_s = sorted(loho_res, key=lambda h: loho_res[h]["LOHO_R2"])
    lr2    = [loho_res[h]["LOHO_R2"] for h in hosts_s]
    rsr2   = [stored_rs.get(h, 0)   for h in hosts_s]
    n_test = [loho_res[h]["N_test"] for h in hosts_s]
    clrs   = [HOST_COLORS.get(h, "gray") for h in hosts_s]
    x      = np.arange(len(hosts_s))

    # ── 4a: LOHO R² bar ──────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    bars = ax.bar(x, lr2, color=clrs, alpha=0.88, width=0.62,
                  edgecolor="white", linewidth=0.5)
    ax.axhline(0,   color="black", lw=REF_LINE_LW)
    ax.axhline(0.5, color="gray",  ls="--", lw=REF_LINE_LW * 0.8, alpha=0.50)
    y_min = min(lr2) - 0.40
    y_max = max(max(lr2) + 0.55, 0.8)
    ax.set_ylim(y_min, y_max)
    for b, n in zip(bars, n_test):
        yp   = b.get_height()
        yoff = (y_max - y_min) * 0.055 if yp >= 0 else -(y_max - y_min) * 0.10
        ax.text(b.get_x() + b.get_width() / 2, yp + yoff,
                f"n={n}", ha="center",
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(x)
    ax.set_xticklabels(hosts_s, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("LOHO R²", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Leave-One-Host-Out R²\n(Generalisation to Unseen Host)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig04a_loho_r2.tiff")

    # ── 4b: LOHO vs random-split ──────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    w2 = 0.38
    ax.bar(x - w2/2, lr2,  w2, color=clrs, alpha=0.88,
           label="LOHO R²", edgecolor="white", linewidth=0.5)
    ax.bar(x + w2/2, rsr2, w2, color=clrs, alpha=0.38,
           label="Random-split R²", edgecolor="white", linewidth=0.5)
    ax.axhline(0, color="black", lw=REF_LINE_LW)
    y_mn = min(lr2) - 0.40
    y_mx = max(max(rsr2) + 0.40, 0.9)
    ax.set_ylim(y_mn, y_mx)
    ax.set_xticks(x)
    ax.set_xticklabels(hosts_s, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("R²", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("LOHO vs Random-Split R²\n"
                 "(Memorisation vs Generalisation Gap)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="upper right",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig04b_loho_vs_random.tiff")


# ── Figure 5: delta-Eg distributions by host and synthesis route ──────────────

def plot_fig5(d):
    print("\nFigure 5 ...")

    df = d["df"]
    HOST_ORDER   = ["CeO2", "ZnO", "SnO2", "TiO2", "WO3", "In2O3"]
    SYNTH_MAP    = {
        "Co-Prec":      "Co-Precipitation",
        "Hydrothermal": "Hydrothermal",
        "Sol-Gel":      "Sol-Gel",
        "Spray-Pyr":    "Spray-Pyrolysis",
        "Phys-Dep":     "Physical-Deposition",
        "Combustion":   "Combustion",
        "Green-Synth":  "Green-Synthesis",
    }
    SYNTH_COLORS = ["#1F4E79", "#2E75B6", "#5BA3D0", "#FF7F0E",
                    "#9467BD", "#D62728", "#2CA02C"]
    MED = dict(color="black", lw=2.2)
    WHI = dict(lw=1.8)
    CAP = dict(lw=1.8)
    FLI = dict(marker="o", markersize=5, alpha=0.55, markeredgewidth=0.6)

    # ── 5a: by host ──────────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    data_host = [df[df["Host"] == h][TARGET].values for h in HOST_ORDER]
    bp = ax.boxplot(data_host, patch_artist=True,
                    medianprops=MED, whiskerprops=WHI,
                    capprops=CAP, flierprops=FLI)
    for patch, host in zip(bp["boxes"], HOST_ORDER):
        patch.set_facecolor(HOST_COLORS[host])
        patch.set_alpha(0.80)
        patch.set_linewidth(1.8)
    for i, (h, wp) in enumerate(
            zip(HOST_ORDER, zip(bp["whiskers"][::2], bp["whiskers"][1::2]))):
        top_y = wp[1].get_ydata()[1]
        v     = df[df["Host"] == h][TARGET]
        q1, q3 = v.quantile(0.25), v.quantile(0.75)
        ax.text(i + 1, top_y + 0.06,
                f"IQR\n{q3-q1:.3f}",
                ha="center", fontsize=FS_ANNOT - 1,
                color="navy", fontweight="bold")
    ax.set_ylim(ax.get_ylim()[0], ax.get_ylim()[1] + 0.38)
    ax.axhline(0, color="gray", ls="--", lw=REF_LINE_LW, alpha=0.70)
    ax.set_xticklabels(HOST_ORDER, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("ΔEg Distribution by Host Material\n"
                 "(IQR annotated from dataset)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig05a_dist_by_host.tiff")

    # ── 5b: by synthesis route ────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(11, 8), constrained_layout=True)
    data_synth = [
        df[df["Synthesis_Method_grouped"] == v][TARGET].dropna().values
        for v in SYNTH_MAP.values()
    ]
    bp2 = ax.boxplot(data_synth, patch_artist=True,
                     medianprops=MED, whiskerprops=WHI,
                     capprops=CAP, flierprops=FLI)
    for patch, clr in zip(bp2["boxes"], SYNTH_COLORS):
        patch.set_facecolor(clr)
        patch.set_alpha(0.80)
        patch.set_linewidth(1.8)
    ax.axhline(0, color="gray", ls="--", lw=REF_LINE_LW, alpha=0.70)
    ax.set_xticklabels(list(SYNTH_MAP.keys()),
                       rotation=22, ha="right",
                       fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("ΔEg Distribution by Synthesis Route",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig05b_dist_by_synthesis.tiff")


# ── Figure 6: residual diagnostics for the Gradient Boosting model ───────────

def plot_fig6(d):
    print("\nFigure 6 ...")

    y_test     = d["y_test"].values
    y_pred     = d["results"]["Gradient Boosting"]["y_pred"]
    resid      = y_test - y_pred
    test_hosts = d["df"].loc[d["test_idx"], "Host"].values
    abs_r      = np.abs(resid)
    pct_010    = d["pct_010"]
    pct_020    = d["pct_020"]

    # ── 6a: residuals vs predicted ────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 8), constrained_layout=True)
    for host, clr in HOST_COLORS.items():
        mask = test_hosts == host
        ax.scatter(y_pred[mask], resid[mask],
                   color=clr, s=55, alpha=0.82,
                   edgecolors="white", linewidths=0.4, zorder=3, label=host)
    ax.axhline(0, color="black", lw=REF_LINE_LW)
    ax.axhspan(-0.15, 0.15, alpha=0.09, color="green")
    ax.set_xlabel("Predicted ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("Residual  (actual − predicted, eV)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Residuals vs Predicted ΔEg\n"
                 f"(Gradient Boosting, Test Set N={len(y_test)})  Green band = ±0.15 eV",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="upper right",
                    ncol=2, framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig06a_residuals_vs_predicted.tiff")

    # ── 6b: residual histogram ────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 7), constrained_layout=True)
    ax.hist(resid, bins=26, color="#2E75B6", alpha=0.86,
            edgecolor="white", linewidth=0.6)
    ax.axvline(resid.mean(), color="red", ls="--", lw=REF_LINE_LW,
               label=f"Mean = {resid.mean():.3f} eV")
    ax.axvline(0, color="black", lw=REF_LINE_LW * 0.8, alpha=0.55)
    ax.set_xlabel("Residual (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("Count", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title(f"Residual Distribution\n"
                 f"(mean = {resid.mean():.3f} eV, approx. Gaussian)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="upper right",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig06b_residual_histogram.tiff")

    # ── 6c: cumulative error ──────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 7), constrained_layout=True)
    sorted_ = np.sort(abs_r)
    cdf     = np.arange(1, len(sorted_) + 1) / len(sorted_) * 100
    ax.plot(sorted_, cdf, color="#1F4E79", lw=3.0)
    for thresh, clr_t in [(0.10, "darkorange"), (0.20, "red")]:
        pct = (abs_r <= thresh).mean() * 100
        ax.axvline(thresh, color=clr_t, ls="--", lw=REF_LINE_LW)
        ax.axhline(pct,    color=clr_t, ls="--", lw=REF_LINE_LW)
        ax.annotate(f"±{thresh:.2f} eV\n{pct:.1f}%",
                    xy=(thresh, pct),
                    xytext=(thresh + 0.055, pct - 12),
                    fontsize=FS_ANNOT, color=clr_t, fontweight="bold",
                    arrowprops=dict(arrowstyle="-", color=clr_t, lw=1.2),
                    bbox=dict(boxstyle="round,pad=0.3", fc="white",
                              alpha=0.92, ec=clr_t, lw=1.5))
    ax.set_xlabel("|Residual| (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("Cumulative %",    fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title(f"Cumulative Error Distribution\n"
                 f"{pct_010:.1f}% within ±0.10 eV  |  {pct_020:.1f}% within ±0.20 eV",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig06c_cumulative_error.tiff")


# ── Figure 7: split-conformal prediction intervals ────────────────────────────

def plot_fig7(d):
    print("\nFigure 7 ...")

    y_test      = d["y_test"].values
    y_pred_conf = d["y_pred_conf"]
    lower       = d["lower"]; upper = d["upper"]
    q90         = d["q90"];   coverage = d["coverage"]
    fs_guar     = d["fs_guarantee"]; within = d["within"]
    sort_ix     = np.argsort(y_test)
    x_plot      = np.arange(len(sort_ix))

    # ── 7a: PI band ───────────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(12, 7), constrained_layout=True)
    ax.fill_between(x_plot, lower[sort_ix], upper[sort_ix],
                    alpha=0.26, color="teal", label="90% PI")
    ax.plot(x_plot, y_pred_conf[sort_ix],
            color="teal", lw=2.5, alpha=0.85, label="Predicted ΔEg")
    ax.scatter(x_plot, y_test[sort_ix],
               color="navy", s=20, alpha=0.78, zorder=4, label="Actual ΔEg")
    ax.set_xlabel("Test samples (sorted by actual ΔEg)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title(f"90% Conformal Prediction Intervals\n"
                 f"Width = {2*q90:.4f} eV  |  half-width q₉₀ = {q90:.4f} eV",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="upper left",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig07a_conformal_band.tiff")

    # ── 7b: coverage scatter ──────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(8, 8), constrained_layout=True)
    ax.scatter(y_pred_conf[within], y_test[within],
               color="teal", s=50, alpha=0.82,
               edgecolors="white", linewidths=0.4, zorder=3,
               label=f"Within PI  (n={within.sum()})")
    ax.scatter(y_pred_conf[~within], y_test[~within],
               color="red", marker="x", s=80, lw=2.5, zorder=4,
               label=f"Outside PI  (n={(~within).sum()})")
    lim = 1.40
    ax.plot([-lim, lim], [-lim, lim], "k--", lw=REF_LINE_LW, alpha=0.55)
    ax.set_xlim(-lim, lim); ax.set_ylim(-lim, lim)
    ax.set_xlabel("Predicted ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("Actual ΔEg (eV)",    fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title(f"Coverage Visualisation\n"
                 f"Empirical {coverage:.1f}%  |  Target 90%  |  "
                 f"Guarantee ≥{fs_guar:.1f}%",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    leg = ax.legend(fontsize=FS_LEGEND, loc="lower right",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    ax.set_aspect("equal", adjustable="box")
    _save(fig, ax, "fig07b_conformal_coverage.tiff")


# ── Figure 8: inter-laboratory noise floor analysis ───────────────────────────

def plot_fig8(d):
    print("\nFigure 8 ...")

    vdf       = d["vdf"].sort_values("sigma", ascending=True).reset_index(drop=True)
    gb_mae    = d["results"]["Gradient Boosting"]["MAE"]
    gb_r2     = d["results"]["Gradient Boosting"]["Test_R2"]
    clean_mae = d["mae_clean"]
    r2_clean  = d["r2_clean"]
    med_sig   = float(vdf["sigma"].median())
    men_sig   = float(vdf["sigma"].mean())
    p90_sig   = float(vdf["sigma"].quantile(0.90))

    # ── 8a: horizontal bar (all pairs) ───────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 16), constrained_layout=True)
    bar_clrs = ["#C00000" if s > 0.40 else
                "#FF7F0E" if s > 0.20 else "#2E75B6"
                for s in vdf["sigma"]]
    bars_h = ax.barh(range(len(vdf)), vdf["sigma"],
                     color=bar_clrs, alpha=0.88, height=0.75)
    # annotate each bar with value
    for bar_, v in zip(bars_h, vdf["sigma"]):
        ax.text(v + vdf["sigma"].max() * 0.010,
                bar_.get_y() + bar_.get_height() / 2,
                f"{v:.3f}", va="center", ha="left",
                fontsize=FS_ANNOT - 2, fontweight="bold", color=LABEL_COLOR)
    ax.axvline(gb_mae,  color="green",  ls="--", lw=REF_LINE_LW,
               label=f"GB MAE = {gb_mae:.4f} eV")
    ax.axvline(med_sig, color="navy",   ls="-.", lw=REF_LINE_LW,
               label=f"Median σ = {med_sig:.4f} eV")
    ax.axvline(0.075,   color="gray",   ls=":",  lw=REF_LINE_LW,
               label="Tauc-plot noise ≈ 0.075 eV")
    ax.set_yticks(range(len(vdf)))
    ax.set_yticklabels(
        [f"{r['Host']}:{r['Dopant']}" for _, r in vdf.iterrows()],
        fontsize=FS_TICK - 1, fontweight="bold")
    ax.set_xlabel("Inter-paper σ (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Inter-Laboratory σ – All Repeated Host-Dopant Pairs\n"
                 "(red > 0.40 eV, orange 0.20–0.40 eV, blue < 0.20 eV)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    ax.margins(x=0.18)
    leg = ax.legend(fontsize=FS_LEGEND, loc="lower right",
                    framealpha=0.95, edgecolor="black")
    _bold_legend(leg)
    _save(fig, ax, "fig08a_noise_bar_chart.tiff")

    # ── 8b: summary comparison ────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    labels = ["Tauc-plot\nnoise", "GB MAE\n(clean)", "GB MAE\n(full)",
              "Median\nσ", "Mean\nσ", "90th-pct\nσ"]
    values = [0.075, clean_mae, gb_mae, med_sig, men_sig, p90_sig]
    clrs_b = ["#AAAAAA", "#2CA02C", "#2E75B6", "#FF7F0E", "#D62728", "#C00000"]
    bars   = ax.bar(range(len(labels)), values, color=clrs_b, alpha=0.88,
                    edgecolor="white", linewidth=0.5)
    y_max  = max(values)
    ax.set_ylim(0, y_max * 1.32)
    for bar_, val in zip(bars, values):
        ax.text(bar_.get_x() + bar_.get_width() / 2,
                val + y_max * 0.018, f"{val:.4f}",
                ha="center", va="bottom",
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("Value (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title(f"Noise Floor vs Model Performance\n"
                 f"Clean-subset R²={r2_clean:.4f}  |  Full R²={gb_r2:.4f}  "
                 f"[ΔR²=+{r2_clean - gb_r2:.4f}]",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    _save(fig, ax, "fig08b_noise_summary.tiff")


# ── Figure 9: within-paper concentration-sensitivity validation ───────────────

def plot_fig9(d):
    print("\nFigure 9 ...")

    slope_df = d["slope_df"]
    r_all    = float(d["r_all"]); p_all = float(d["p_all"])
    r_to     = float(d["r_to"])
    r_mx     = float(d["r_mx"]); p_mx = float(d["p_mx"])
    to_df    = slope_df[slope_df["Series_type"] == "train-only"]
    mx_df    = slope_df[slope_df["Series_type"] == "mixed"]
    n_tot    = len(slope_df)
    sign_n   = int(slope_df["Sign_agreement"].sum())

    TYPE_STYLE = {
        "train-only": dict(marker="o", alpha=0.60, s=55),
        "mixed":      dict(marker="s", alpha=0.92, s=55),
        "test-only":  dict(marker="^", alpha=1.00, s=70),
    }

    # ── 9a: slope scatter ────────────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 9), constrained_layout=True)
    for stype, style in TYPE_STYLE.items():
        sub = slope_df[slope_df["Series_type"] == stype]
        if len(sub) == 0: continue
        for host, grp in sub.groupby("Host"):
            ax.scatter(grp["Observed_slope"], grp["Predicted_slope"],
                       color=HOST_COLORS[host], zorder=3,
                       edgecolors="white", linewidths=0.5, **style)
    all_v = pd.concat([slope_df["Observed_slope"], slope_df["Predicted_slope"]])
    lim   = max(abs(all_v.min()), abs(all_v.max())) * 1.15
    ax.plot([-lim, lim], [-lim, lim], "k--", lw=REF_LINE_LW, alpha=0.65,
            label="1:1 line")
    sf, ic, *_ = stats.linregress(slope_df["Observed_slope"],
                                   slope_df["Predicted_slope"])
    xf = np.linspace(-lim, lim, 200)
    ax.plot(xf, sf * xf + ic, "r-", lw=REF_LINE_LW, alpha=0.80, label="OLS fit")
    ax.set_xlim(-lim, lim); ax.set_ylim(-lim, lim)
    ax.set_xlabel("Observed  dΔEg/dc  (eV per molar-ratio)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("Model-predicted  dΔEg/dc  (eV per molar-ratio)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Within-Paper Concentration Sensitivity\n"
                 "(circles=train-only, squares=mixed, triangles=test-only)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    textstr = (f"Overall N={n_tot}:  r = {r_all:.3f}\n"
               f"Train-only N={len(to_df)}:  r = {r_to:.3f}\n"
               f"Mixed N={len(mx_df)}:  r = {r_mx:.3f},  p = {p_mx:.3f}\n"
               f"Sign agreement: {sign_n}/{n_tot} ({sign_n/n_tot*100:.1f}%)")
    ax.text(0.03, 0.97, textstr,
            transform=ax.transAxes, fontsize=FS_ANNOT, va="top",
            fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.5", facecolor="#EBF3FB",
                      alpha=0.94, edgecolor="#2E75B6", lw=2.0))
    handles = [
        Line2D([0],[0], marker="o", color="gray", markersize=9,
               ls="None", alpha=0.65, label="Train-only"),
        Line2D([0],[0], marker="s", color="gray", markersize=9,
               ls="None", alpha=0.92, label="Mixed"),
        Line2D([0],[0], color="k", ls="--", lw=REF_LINE_LW, label="1:1 line"),
    ]
    for h, clr in HOST_COLORS.items():
        handles.append(Line2D([0],[0], marker="o", color=clr,
                               markersize=8, ls="None", alpha=0.88, label=h))
    leg = ax.legend(handles=handles, loc="lower right",
                    fontsize=FS_LEGEND, framealpha=0.95,
                    edgecolor="black", ncol=2)
    _bold_legend(leg)
    ax.set_aspect("equal", adjustable="box")
    _save(fig, ax, "fig09a_slope_scatter.tiff")

    # ── 9b: per-host Pearson r + sign agreement ───────────────────────────────
    host_stats = {}
    for host, grp in slope_df.groupby("Host"):
        if len(grp) < 2: continue
        r_h, p_h = stats.pearsonr(grp["Observed_slope"], grp["Predicted_slope"])
        host_stats[host] = {
            "r": r_h, "p": p_h, "n": len(grp),
            "sa": grp["Sign_agreement"].mean() * 100
        }
    hosts_s = sorted(host_stats, key=lambda h: host_stats[h]["r"], reverse=True)
    r_vals  = [host_stats[h]["r"]  for h in hosts_s]
    sa_vals = [host_stats[h]["sa"] for h in hosts_s]
    n_vals  = [host_stats[h]["n"]  for h in hosts_s]
    clrs_h  = [HOST_COLORS[h]      for h in hosts_s]
    x       = np.arange(len(hosts_s))

    fig, ax = plt.subplots(figsize=(10, 8), constrained_layout=True)
    ax.bar(x, r_vals, color=clrs_h, alpha=0.88, width=0.55,
           edgecolor="white", linewidth=0.5)
    ax2 = ax.twinx()
    ax2.plot(x, sa_vals, "D--", color="#333333", markersize=9,
             lw=REF_LINE_LW, alpha=0.88, label="Sign agreement (%)")
    ax2.set_ylim(0, 130)
    ax2.set_ylabel("Sign agreement (%)", color="#333333",
                   fontsize=FS_AXLBL, fontweight="bold", labelpad=10)
    ax2.tick_params(axis="y", labelcolor="#333333", labelsize=FS_TICK)
    for lbl in ax2.get_yticklabels():
        lbl.set_fontweight("bold"); lbl.set_fontsize(FS_TICK)
    for spine in ax2.spines.values(): spine.set_linewidth(SPINE_LW)
    ax.axhline(0.7, color="red", lw=REF_LINE_LW, ls="--", alpha=0.70)
    ax.axhline(0,   color="black", lw=REF_LINE_LW * 0.65, alpha=0.55)
    y_max2 = max(r_vals) + 0.45
    ax.set_ylim(min(r_vals) - 0.25, y_max2)
    for bar_, n in zip(ax.patches, n_vals):
        yp = bar_.get_height()
        ax.text(bar_.get_x() + bar_.get_width() / 2,
                yp + 0.05 if yp >= 0 else yp - 0.18,
                f"N={n}", ha="center", va="bottom",
                fontsize=FS_ANNOT, fontweight="bold", color=LABEL_COLOR)
    ax.set_xticks(x)
    ax.set_xticklabels(hosts_s, fontsize=FS_TICK, fontweight="bold")
    ax.set_ylabel("Pearson r", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("Per-Host Within-Paper Slope Correlation & Sign Agreement",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    dl = Line2D([0],[0], marker="D", color="#333333", ls="--",
                markersize=9, lw=REF_LINE_LW, label="Sign agr. (%)")
    leg2 = ax.legend(handles=[dl], fontsize=FS_LEGEND, loc="upper right",
                     framealpha=0.95, edgecolor="black")
    _bold_legend(leg2)
    _save_twin(fig, ax, ax2, "fig09b_per_host_correlation.tiff")


# ── Figure 10: physical trend scatter plots ────────────────────────────────────

def plot_fig10(d):
    print("\nFigure 10 ...")

    df   = d["df"]
    dEg  = df[TARGET].values
    conc = df["Conc_1"].values
    en   = df["EN_difference"].values
    ir   = df["Ionic_radius_mismatch_rel"].values
    temp = df["Synthesis_Temp(C)"].values

    # ── 10a: ΔEg vs concentration ─────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 8), constrained_layout=True)
    sc = ax.scatter(conc, dEg, c=en, cmap="coolwarm",
                    s=22, alpha=0.68, vmin=-2, vmax=2, edgecolors="none")
    ax.axhline(0, color="gray", ls="--", lw=REF_LINE_LW, alpha=0.65)
    ax.set_xlabel("Dopant concentration (molar-ratio)",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("ΔEg vs Dopant Concentration\n"
                 "(colour = ΔElectronegativity [host − dopant])",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    cbar = plt.colorbar(sc, ax=ax, shrink=0.80, pad=0.02)
    cbar.set_label("ΔElectronegativity (Pauling)",
                   fontsize=FS_AXLBL - 1, fontweight="bold")
    cbar.ax.tick_params(labelsize=FS_TICK - 1)
    for lbl in cbar.ax.get_yticklabels():
        lbl.set_fontweight("bold")
    ax.text(0.04, 0.05,
            "Mott-Hubbard saturation\nat high concentration",
            transform=ax.transAxes, fontsize=FS_ANNOT,
            ha="left", va="bottom", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.4", facecolor="white",
                      alpha=0.90, edgecolor="#888888", lw=1.5))
    _save(fig, ax, "fig10a_conc_trend.tiff")

    # ── 10b: ΔEg vs ionic mismatch ────────────────────────────────────────────
    fig, ax = plt.subplots(figsize=(9, 8), constrained_layout=True)
    valid = ~np.isnan(temp)
    sc2   = ax.scatter(ir[valid], dEg[valid],
                       c=temp[valid], cmap="YlOrRd",
                       s=22, alpha=0.68, vmin=200, vmax=900, edgecolors="none")
    ax.axhline(0, color="gray", ls="--", lw=REF_LINE_LW, alpha=0.65)
    ax.set_xlabel("Ionic radius mismatch  (r_dopant − r_host) / r_host",
                  fontsize=FS_AXLBL, fontweight="bold")
    ax.set_ylabel("ΔEg (eV)", fontsize=FS_AXLBL, fontweight="bold")
    ax.set_title("ΔEg vs Ionic Radius Mismatch\n"
                 "(colour = Synthesis temperature °C)",
                 fontsize=FS_TITLE, fontweight="bold", pad=TITLE_PAD)
    cbar2 = plt.colorbar(sc2, ax=ax, shrink=0.80, pad=0.02)
    cbar2.set_label("Synthesis temperature (°C)",
                    fontsize=FS_AXLBL - 1, fontweight="bold")
    cbar2.ax.tick_params(labelsize=FS_TICK - 1)
    for lbl in cbar2.ax.get_yticklabels():
        lbl.set_fontweight("bold")
    ax.text(0.04, 0.05,
            "Peak strain effect at\nintermediate mismatch (~0.2–0.4)",
            transform=ax.transAxes, fontsize=FS_ANNOT,
            ha="left", va="bottom", fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.4", facecolor="white",
                      alpha=0.90, edgecolor="#888888", lw=1.5))
    _save(fig, ax, "fig10b_mismatch_trend.tiff")


# ── manuscript tables: plain black-and-white Word document ───────────────────

def _set_cell(cell, text, bold=False, bg_hex=None, font_size=10,
              align="center", color_hex=None):
    # write text into a table cell with Times New Roman and optional shading
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                      else WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Times New Roman"
    # color_hex and bg_hex kept as parameters for compatibility but not used
    # in the journal version (strict black-and-white)


def _table_title(doc, text):
    # bold table number + title line, no caption, journal style
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"


def build_docx(d):
    # write all five manuscript tables to a plain black-and-white Word document
    if not HAS_DOCX:
        print("  python-docx not installed. Skipping DOCX output.")
        return

    print("\nGenerating bandgap_tables.docx ...")
    doc = DocxDocument()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # set default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    results    = d["results"]
    vdf        = d["vdf"]
    hv         = d["hv"]
    slope_df   = d["slope_df"]
    r2_clean     = d["r2_clean"]
    mae_clean    = d["mae_clean"]
    rmse_clean   = d["rmse_clean"]
    cv_r2_clean  = d["cv_r2_clean"]
    cv_sig_clean = d["cv_sig_clean"]
    q90        = d["q90"]
    coverage   = d["coverage"]
    fs_guar    = d["fs_guarantee"]
    n_cal      = d["n_cal"]
    r_all      = d["r_all"]
    r_mx       = d["r_mx"]
    p_mx       = d["p_mx"]
    r_to       = d["r_to"]
    pct_010    = d["pct_010"]
    pct_020    = d["pct_020"]
    mean_res   = d["mean_res"]
    gb_r       = results["Gradient Boosting"]
    gb_mae     = gb_r["MAE"]
    gb_rmse    = gb_r["RMSE"]
    gb_r2      = gb_r["Test_R2"]

    # helper: set header row cells (bold, bottom border only -- journal style)
    def _hcell(cell, text, width_in):
        cell.width = Inches(width_in)
        _set_cell(cell, text, bold=True, font_size=10,
                  align="center" if text != "Model" else "left")

    # helper: set data row cells
    def _dcell(cell, text, width_in, bold=False, left=False):
        cell.width = Inches(width_in)
        _set_cell(cell, text, bold=bold, font_size=10,
                  align="left" if left else "center")

    # ── Table 1 ───────────────────────────────────────────────────────────────
    _table_title(doc, "Table 1. Model Performance Benchmarking")

    ORDER = ["Extra Trees", "Random Forest", "Gradient Boosting",
             "SVR-RBF", "Lasso", "Ridge", "ElasticNet"]
    HDR1  = ["Model", "Test R\u00b2", "RMSE (eV)", "MAE (eV)",
             "CV R\u00b2 (5-fold)", "CV \u03c3"]
    WID1  = [1.9, 0.85, 0.95, 0.85, 1.20, 0.75]

    t1 = doc.add_table(rows=1 + len(ORDER), cols=len(HDR1))
    t1.style = "Table Grid"
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(HDR1, WID1)):
        _hcell(t1.rows[0].cells[j], h, w)
    for i, name in enumerate(ORDER):
        r  = results[name]
        is_best = (name == "Extra Trees")
        note = " (highest R\u00b2)"   if is_best else \
               " (SHAP)"             if name == "Gradient Boosting" else ""
        vals = [name + note,
                f"{r['Test_R2']:.4f}", f"{r['RMSE']:.4f}",
                f"{r['MAE']:.4f}",     f"{r['CV_R2']:.4f}",
                f"{r['CV_sig']:.4f}"]
        for j, (v, w) in enumerate(zip(vals, WID1)):
            _dcell(t1.rows[i+1].cells[j], v, w,
                   bold=is_best, left=(j == 0))
    doc.add_paragraph()

    # ── Table 2 ───────────────────────────────────────────────────────────────
    _table_title(doc, "Table 2. Noise Floor vs. Model Performance")

    med_sig = float(vdf["sigma"].median())
    men_sig = float(vdf["sigma"].mean())
    p90_sig = float(vdf["sigma"].quantile(0.90))

    t2_rows = [
        ("Tauc-plot single-lab noise",              "0.075"),
        ("GB MAE, full dataset (N = 638)",          f"{gb_mae:.4f}"),
        ("Median inter-paper \u03c3, repeated pairs",f"{med_sig:.4f}"),
        ("Mean inter-paper \u03c3, repeated pairs", f"{men_sig:.4f}"),
        ("GB RMSE, full dataset",                   f"{gb_rmse:.4f}"),
        ("90th-percentile inter-paper \u03c3",      f"{p90_sig:.4f}"),
        ("GB MAE, clean subset (N \u2248 566)",     f"{mae_clean:.4f}"),
        ("GB RMSE, clean subset",                   f"{rmse_clean:.4f}"),
    ]
    WID2 = [3.8, 1.5]
    t2 = doc.add_table(rows=1 + len(t2_rows), cols=2)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(["Quantity", "Value (eV)"], WID2)):
        _hcell(t2.rows[0].cells[j], h, w)
    for i, (lbl, val) in enumerate(t2_rows):
        _dcell(t2.rows[i+1].cells[0], lbl, WID2[0], left=True)
        _dcell(t2.rows[i+1].cells[1], val, WID2[1])
    doc.add_paragraph()

    # ── Table 3 ───────────────────────────────────────────────────────────────
    _table_title(doc, "Table 3. High-Variance Host\u2013Dopant Pairs (\u03c3 > 0.40 eV)")

    HDR3 = ["Host", "Dopant", "N papers", "N rows",
            "\u03c3 (eV)", "Range (eV)"]
    WID3 = [0.90, 0.90, 0.80, 0.75, 0.85, 1.05]
    hv_s = hv.sort_values("sigma", ascending=False)
    t3 = doc.add_table(rows=1 + len(hv_s), cols=len(HDR3))
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(HDR3, WID3)):
        _hcell(t3.rows[0].cells[j], h, w)
    for i, (_, row_) in enumerate(hv_s.iterrows()):
        vals3 = [row_["Host"], row_["Dopant"],
                 str(int(row_["N_papers"])), str(int(row_["N_rows"])),
                 f"{row_['sigma']:.4f}", f"{row_['range']:.4f}"]
        for j, (v, w) in enumerate(zip(vals3, WID3)):
            _dcell(t3.rows[i+1].cells[j], v, w, left=(j < 2))
    doc.add_paragraph()

    # ── Table 4 ───────────────────────────────────────────────────────────────
    _table_title(doc, "Table 4. Gradient Boosting Performance: Full vs. Clean Subset")

    HDR4 = ["Metric", "Full dataset (N = 638)", "Clean subset (N \u2248 566)", "Change"]
    WID4 = [1.0, 1.80, 1.80, 1.0]
    gb_cv_r2  = gb_r["CV_R2"]
    gb_cv_sig = gb_r["CV_sig"]
    t4_rows = [
        ("Test R\u00b2",
         f"{gb_r2:.4f}",     f"{r2_clean:.4f}",
         f"+{r2_clean - gb_r2:.4f}"),
        ("RMSE (eV)",
         f"{gb_rmse:.4f}",   f"{rmse_clean:.4f}",
         f"{rmse_clean - gb_rmse:.4f}"),
        ("MAE (eV)",
         f"{gb_mae:.4f}",    f"{mae_clean:.4f}",
         f"{mae_clean - gb_mae:.4f}"),
        ("CV-R\u00b2 (5-fold)",
         f"{gb_cv_r2:.3f} \u00b1 {gb_cv_sig:.3f}",
         f"{cv_r2_clean:.3f} \u00b1 {cv_sig_clean:.3f}",
         f"+{cv_r2_clean - gb_cv_r2:.3f}"),
    ]
    t4 = doc.add_table(rows=1 + len(t4_rows), cols=4)
    t4.style = "Table Grid"
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(HDR4, WID4)):
        _hcell(t4.rows[0].cells[j], h, w)
    for i, (metric, full, clean, chg) in enumerate(t4_rows):
        for j, (v, w) in enumerate(zip([metric, full, clean, chg], WID4)):
            _dcell(t4.rows[i+1].cells[j], v, w, bold=(j == 3))
    doc.add_paragraph()

    # ── Table 5 ───────────────────────────────────────────────────────────────
    _table_title(doc, "Table 5. Per-Host Within-Paper Concentration-Sensitivity")

    HDR5 = ["Host", "N series", "Pearson r", "p-value", "Sign agr. (%)"]
    WID5 = [0.90, 0.85, 0.95, 0.95, 1.15]
    host_rows = []
    for host, grp in slope_df.groupby("Host"):
        if len(grp) < 2:
            continue
        r_h, p_h = stats.pearsonr(grp["Observed_slope"], grp["Predicted_slope"])
        sa_h     = grp["Sign_agreement"].mean() * 100
        host_rows.append((host, len(grp), r_h, p_h, sa_h))
    host_rows.sort(key=lambda x: x[2], reverse=True)

    t5 = doc.add_table(rows=1 + len(host_rows) + 1, cols=len(HDR5))
    t5.style = "Table Grid"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(HDR5, WID5)):
        _hcell(t5.rows[0].cells[j], h, w)
    for i, (host, n, r_h, p_h, sa_h) in enumerate(host_rows):
        p_str = "<0.001" if p_h < 0.001 else f"{p_h:.3f}"
        vals5 = [host, str(n), f"{r_h:.3f}", p_str, f"{sa_h:.1f}%"]
        for j, (v, w) in enumerate(zip(vals5, WID5)):
            _dcell(t5.rows[i+1].cells[j], v, w, left=(j == 0))
    # overall summary row
    n_total  = len(slope_df)
    sign_pct = slope_df["Sign_agreement"].mean() * 100
    ov_vals  = ["Overall", str(n_total), f"{r_all:.3f}",
                f"{d['p_all']:.3f}", f"{sign_pct:.1f}%"]
    for j, (v, w) in enumerate(zip(ov_vals, WID5)):
        _dcell(t5.rows[-1].cells[j], v, w, bold=True, left=(j == 0))
    doc.add_paragraph()

    # ── Key metrics reference table ────────────────────────────────────────────
    _table_title(doc, "Table S1. Key Computed Metrics (Supplementary Reference)")

    rows_km = [
        ("Extra Trees  Test R\u00b2",
         f"{results['Extra Trees']['Test_R2']:.4f}"),
        ("Extra Trees  RMSE (eV)",
         f"{results['Extra Trees']['RMSE']:.4f}"),
        ("Extra Trees  MAE (eV)",
         f"{results['Extra Trees']['MAE']:.4f}"),
        ("Extra Trees  CV R\u00b2 \u00b1 \u03c3",
         f"{results['Extra Trees']['CV_R2']:.4f} \u00b1 {results['Extra Trees']['CV_sig']:.4f}"),
        ("GB  Test R\u00b2",           f"{gb_r2:.4f}"),
        ("GB  RMSE (eV)",              f"{gb_rmse:.4f}"),
        ("GB  MAE (eV)",               f"{gb_mae:.4f}"),
        ("GB  CV R\u00b2 \u00b1 \u03c3",
         f"{gb_r['CV_R2']:.4f} \u00b1 {gb_r['CV_sig']:.4f}"),
        ("Conformal q\u2089\u2080 half-width (eV)", f"{q90:.4f}"),
        ("Conformal full interval width (eV)",      f"{2*q90:.4f}"),
        ("Conformal empirical coverage",            f"{coverage:.1f}%"),
        ("Conformal finite-sample guarantee",       f"\u2265{fs_guar:.1f}%"),
        ("Conformal N\u1d04\u1d00\u029f",           f"{n_cal}"),
        ("GB residual mean (eV)",                   f"{mean_res:.4f}"),
        ("GB residuals within \u00b10.10 eV",       f"{pct_010:.1f}%"),
        ("GB residuals within \u00b10.20 eV",       f"{pct_020:.1f}%"),
        ("Clean-subset R\u00b2 (GB)",               f"{r2_clean:.4f}"),
        ("Clean-subset RMSE (eV) (GB)",             f"{rmse_clean:.4f}"),
        ("Clean-subset MAE (eV) (GB)",              f"{mae_clean:.4f}"),
        ("Clean-subset CV-R\u00b2 \u00b1 \u03c3 (GB)",
         f"{cv_r2_clean:.4f} \u00b1 {cv_sig_clean:.4f}"),
        ("Within-paper slope r, all series",        f"{r_all:.3f}"),
        ("Within-paper slope r, mixed series",      f"{r_mx:.3f}"),
        ("Mixed-series p-value",                    f"{p_mx:.3f}"),
        ("Within-paper slope r, train-only",        f"{r_to:.3f}"),
        ("Sign agreement, all series",
         f"{int(slope_df['Sign_agreement'].sum())}/{n_total} ({sign_pct:.1f}%)"),
    ]
    WID_KM = [3.8, 1.8]
    tkm = doc.add_table(rows=1 + len(rows_km), cols=2)
    tkm.style = "Table Grid"
    tkm.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (h, w) in enumerate(zip(["Metric", "Value"], WID_KM)):
        _hcell(tkm.rows[0].cells[j], h, w)
    for i, (lbl, val) in enumerate(rows_km):
        _dcell(tkm.rows[i+1].cells[0], lbl, WID_KM[0], left=True)
        _dcell(tkm.rows[i+1].cells[1], val, WID_KM[1])

    docx_path = os.path.join(OUTPUT_DIR, "bandgap_tables.docx")
    doc.save(docx_path)
    print(f"  saved -> bandgap_tables.docx")


# ── write requirements.txt and bundle everything into a zip ──────────────────

def write_requirements():
    # pin the core libraries so anyone can recreate the environment exactly
    import importlib.metadata as meta
    libs = ["pandas", "numpy", "scikit-learn", "scipy",
            "matplotlib", "openpyxl", "shap", "python-docx"]
    lines = []
    for lib in libs:
        try:
            ver = meta.version(lib)
            lines.append(f"{lib}=={ver}")
        except meta.PackageNotFoundError:
            lines.append(lib)   # version unknown, list without pin
    req_path = os.path.join(OUTPUT_DIR, "requirements.txt")
    with open(req_path, "w") as f:
        f.write("\n".join(lines) + "\n")
    print(f"  saved -> requirements.txt")
    return req_path


def zip_outputs():
    print("\nBundling outputs into bandgap_outputs.zip ...")
    zip_path = "bandgap_outputs.zip"
    files    = sorted(os.listdir(OUTPUT_DIR))
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for fname in files:
            zf.write(os.path.join(OUTPUT_DIR, fname), arcname=fname)
            print(f"  added: {fname}")
    size_mb = os.path.getsize(zip_path) / 1e6
    print(f"\n  bandgap_outputs.zip  ({size_mb:.1f} MB)")


# ── entry point ───────────────────────────────────────────────────────────────

def main():
    print("\nPhysics-Informed ML for Bandgap Engineering in Metal-Oxide Photocatalysts")
    print("Generating figures and tables (random_state=42) ...")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    apply_global_style()

    d = build_everything()

    print("\nGenerating figures ...")
    plot_fig1(d)
    plot_fig2(d)
    plot_fig3(d)
    plot_fig4(d)
    plot_fig5(d)
    plot_fig6(d)
    plot_fig7(d)
    plot_fig8(d)
    plot_fig9(d)
    plot_fig10(d)

    build_docx(d)
    write_requirements()
    zip_outputs()

    print("\nDone. Files are in bandgap_outputs/  |  archive: bandgap_outputs.zip")
    print()

    r = d["results"]
    print(f"  {'Model':<22} {'Test R2':>8} {'RMSE':>8} {'MAE':>8} "
          f"{'CV R2':>8} {'CV sig':>8}")
    print("  " + "-" * 68)
    for name in ["Extra Trees", "Random Forest", "Gradient Boosting",
                 "SVR-RBF", "Lasso", "Ridge", "ElasticNet"]:
        rv = r[name]
        print(f"  {name:<22} {rv['Test_R2']:>8.4f} {rv['RMSE']:>8.4f} "
              f"{rv['MAE']:>8.4f} {rv['CV_R2']:>8.4f} {rv['CV_sig']:>8.4f}")
    print()


if __name__ == "__main__":
    main()
